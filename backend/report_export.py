"""Generate Word (.docx) downloads of an interview report.

Two flavours:
  - build_full_docx(report)  — everything: scores, criteria table, strengths,
                                gaps, Q&A with model answers, transcript.
  - build_qa_docx(report)    — compact study sheet, just Q + model answer.

We don't introduce any new dependencies — python-docx is already used by the
resume parser.
"""

from datetime import datetime
from io import BytesIO
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


# -----------------------------------------------------------------------------
# Small style helpers
# -----------------------------------------------------------------------------
_VIOLET = RGBColor(0x6B, 0x4D, 0xE0)
_TEXT_DIM = RGBColor(0x55, 0x55, 0x55)
_TEXT_MUTED = RGBColor(0x88, 0x88, 0x88)


def _add_runs(p, *parts):
    """Append (text, {style}) tuples as runs on a paragraph."""
    for item in parts:
        text, style = item if isinstance(item, tuple) else (item, {})
        run = p.add_run(text)
        if style.get("bold"):
            run.bold = True
        if style.get("italic"):
            run.italic = True
        if style.get("color") is not None:
            run.font.color.rgb = style["color"]
        if style.get("size") is not None:
            run.font.size = Pt(style["size"])


def _shade_cell(cell, color_hex: str) -> None:
    """Set a table cell background colour (no API for this in python-docx)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _meta_line(report: dict) -> str:
    bits = []
    if report.get("difficulty"):
        bits.append(report["difficulty"])
    if report.get("candidate_name"):
        bits.append(report["candidate_name"])
    bits.append(datetime.now().strftime("%-d %b %Y"))
    return " · ".join(bits)


def _quote_paragraph(doc: Document, text: str):
    """Indented italic blockquote-style paragraph (works in any Word version)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text or "(no answer given)")
    run.italic = True
    run.font.color.rgb = _TEXT_DIM
    return p


# -----------------------------------------------------------------------------
# Full report
# -----------------------------------------------------------------------------
def build_full_docx(report: dict) -> bytes:
    doc = Document()

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Title ----
    title = doc.add_heading("Interview Report", level=0)
    for r in title.runs:
        r.font.color.rgb = _VIOLET

    # ---- Header info ----
    info = doc.add_paragraph()
    _add_runs(info,
        ("Topic: ", {"bold": True}), f"{report.get('topic', '')}\n",
        ("Difficulty: ", {"bold": True}), f"{report.get('difficulty', '')}\n",
    )
    if report.get("candidate_name"):
        _add_runs(info, ("Candidate: ", {"bold": True}), f"{report['candidate_name']}\n")
    _add_runs(info,
        ("Generated: ", {"bold": True}),
        datetime.now().strftime("%-d %b %Y · %H:%M"),
    )

    # ---- Verdict block ----
    doc.add_paragraph()
    verdict_p = doc.add_paragraph()
    verdict_p.paragraph_format.space_before = Pt(6)
    _add_runs(verdict_p,
        (f"Overall: ", {"bold": True, "size": 14}),
        (f"{report.get('overall_score', '?')}/10  —  ", {"bold": True, "size": 14, "color": _VIOLET}),
        (report.get("verdict", ""), {"bold": True, "size": 14}),
    )

    if report.get("summary"):
        p = doc.add_paragraph(report["summary"])
        p.paragraph_format.space_after = Pt(12)

    # ---- Criteria table ----
    criteria = report.get("criteria") or []
    if criteria:
        doc.add_heading("Criteria", level=1)
        table = doc.add_table(rows=1 + len(criteria), cols=3)
        table.autofit = False
        widths = [Inches(2.0), Inches(0.8), Inches(3.7)]
        for col, w in zip(table.columns, widths):
            for cell in col.cells:
                cell.width = w

        # Header row
        hdr = table.rows[0].cells
        for i, label in enumerate(["Criterion", "Score", "Notes"]):
            hdr[i].text = ""
            run = hdr[i].paragraphs[0].add_run(label)
            run.bold = True
            run.font.color.rgb = _VIOLET
            _shade_cell(hdr[i], "F2EEFF")

        for row_idx, c in enumerate(criteria, 1):
            cells = table.rows[row_idx].cells
            cells[0].text = c.get("name", "")
            cells[1].text = f"{c.get('score', '?')}/10"
            cells[2].text = c.get("rationale", "")
            for cell in cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        doc.add_paragraph()

    # ---- Strengths ----
    strengths = report.get("strengths") or []
    if strengths:
        doc.add_heading("Strengths", level=1)
        for s in strengths:
            doc.add_paragraph(s, style="List Bullet")

    # ---- Gaps ----
    gaps = report.get("gaps") or []
    if gaps:
        doc.add_heading("Gaps & growth areas", level=1)
        for g in gaps:
            doc.add_paragraph(g, style="List Bullet")

    # ---- Questions + model answers ----
    qa = report.get("question_solutions") or []
    if qa:
        doc.add_heading("Questions and model answers", level=1)
        for i, q in enumerate(qa, 1):
            doc.add_heading(f"Q{i}. {q.get('question', '')}", level=2)

            p = doc.add_paragraph()
            _add_runs(p, ("Your answer:", {"bold": True}))
            _quote_paragraph(doc, q.get("candidate_answer") or "")

            p = doc.add_paragraph()
            _add_runs(p, ("Model answer:", {"bold": True}))
            doc.add_paragraph(q.get("model_answer", ""))

            if q.get("feedback"):
                p = doc.add_paragraph()
                _add_runs(p,
                    ("Where you stood: ", {"bold": True, "color": _VIOLET}),
                    (q["feedback"], {"italic": True, "color": _TEXT_DIM}),
                )

    # ---- Transcript ----
    transcript = report.get("transcript") or []
    if transcript:
        doc.add_page_break()
        doc.add_heading("Full transcript", level=1)
        for t in transcript:
            label = "Interviewer" if t.get("role") == "interviewer" else "You"
            p = doc.add_paragraph()
            _add_runs(p,
                (f"{label}: ", {"bold": True, "color": _VIOLET if label == "Interviewer" else _TEXT_DIM}),
                t.get("content", ""),
            )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# -----------------------------------------------------------------------------
# Q&A study sheet
# -----------------------------------------------------------------------------
def build_qa_docx(report: dict) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(f"Q&A — {report.get('topic', 'Interview')}", level=0)
    for r in title.runs:
        r.font.color.rgb = _VIOLET

    meta = doc.add_paragraph()
    run = meta.add_run(_meta_line(report))
    run.italic = True
    run.font.color.rgb = _TEXT_MUTED

    qa = report.get("question_solutions") or []
    if not qa:
        doc.add_paragraph("(No question/answer pairs were captured for this session.)")
    else:
        for i, q in enumerate(qa, 1):
            doc.add_heading(f"Q{i}. {q.get('question', '')}", level=1)

            p = doc.add_paragraph()
            _add_runs(p, ("Your answer:", {"bold": True}))
            _quote_paragraph(doc, q.get("candidate_answer") or "")

            p = doc.add_paragraph()
            _add_runs(p, ("Model answer:", {"bold": True}))
            doc.add_paragraph(q.get("model_answer", ""))

            if q.get("feedback"):
                p = doc.add_paragraph()
                _add_runs(p,
                    ("Where you stood: ", {"bold": True, "color": _VIOLET}),
                    (q["feedback"], {"italic": True, "color": _TEXT_DIM}),
                )
            doc.add_paragraph()  # spacer

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
