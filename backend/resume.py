"""Resume parsing — extract structured candidate profile from PDF/DOCX/TXT.

Two-step pipeline:
  1. Pull plain text out of the file (pypdf / python-docx / read)
  2. Hand the text to Claude with a structured-extraction prompt and
     get back a JSON profile we can drop into the setup form.

The Claude step is what makes this robust — it summarises a noisy
resume into the exact shape the interviewer prompts expect, not raw
bullet-point soup.
"""

import json
import os
import re
from io import BytesIO
from typing import Optional

from claude_agent_sdk import ClaudeAgentOptions, query
from claude_agent_sdk.types import AssistantMessage, TextBlock

MAX_BYTES = 5 * 1024 * 1024   # 5 MB hard cap on uploads
MAX_TEXT_CHARS = 30_000       # truncate very long resumes before LLM


def _model() -> Optional[str]:
    return (
        os.environ.get("CLAUDE_RESUME_MODEL")
        or os.environ.get("CLAUDE_INTERVIEWER_MODEL")
        or os.environ.get("CLAUDE_MODEL")
        or "claude-sonnet-4-6"
    )


# ----- Text extraction -----------------------------------------------------
def extract_text(filename: str, content: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf(content)
    if name.endswith(".docx"):
        return _extract_docx(content)
    if name.endswith(".txt") or name.endswith(".md"):
        return content.decode("utf-8", errors="replace")
    # Fallback: assume text
    try:
        return content.decode("utf-8", errors="replace")
    except Exception:
        raise ValueError(f"Unsupported file type: {filename}")


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader
    try:
        reader = PdfReader(BytesIO(content))
    except Exception as e:
        raise ValueError(f"Could not read PDF: {e}")
    pages = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            pages.append(t.strip())
    text = "\n\n".join(pages).strip()
    if not text:
        raise ValueError(
            "No text extracted from the PDF. It might be an image-only / scanned "
            "resume — convert it to a text PDF or paste your background manually."
        )
    return text


def _extract_docx(content: bytes) -> str:
    from docx import Document
    try:
        doc = Document(BytesIO(content))
    except Exception as e:
        raise ValueError(f"Could not read DOCX: {e}")
    parts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    # Also pull tables (resumes sometimes use them for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("Empty DOCX — no readable text.")
    return text


# ----- Claude extraction ---------------------------------------------------
_EXTRACT_SYSTEM = """You are a resume parser. Given the raw text of a candidate's resume, return a structured profile that an interview-prep tool can use to tailor questions.

Return ONLY a single JSON object, no prose, no markdown fences. Schema:

{
  "name": "<full name, or null if not clearly stated>",
  "headline": "<one short line: current role + years of experience + domain, e.g. 'Senior PM, 6 years in fintech and logistics'>",
  "background": "<3–5 sentence summary written in FIRST PERSON as if the candidate is introducing themselves to an interviewer. Mention current/latest role, 1–2 notable past roles, concrete metrics or achievements, and any distinctive technical or domain depth. Specific > generic.>",
  "role_context_suggestion": "<one short sentence describing what kind of role this person is likely interviewing for next — base it on their trajectory>",
  "focus_areas_suggestion": "<comma-separated list of 4–6 themes that would be relevant to test this candidate on, derived from their experience>",
  "topic_suggestion": "<a short topic string for the interview, e.g. 'Product Management — fintech and payments' or 'Senior backend engineering — distributed systems'>"
}

Rules:
- Pull SPECIFIC details from the resume (real companies, real numbers, real tech). Never invent facts.
- background is the most important field — read like a real person talking, not a CV bullet list.
- If something genuinely isn't in the resume, use null for `name` only; for the others, return your best inference based on what IS in the resume.
- Keep `background` under 600 characters. Keep other fields under 200 characters each.
"""


def _extract_json(text: str) -> dict:
    text = text.strip()
    fence = re.match(r"^```(?:json)?\s*(.*?)\s*```$", text, re.DOTALL)
    if fence:
        text = fence.group(1).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    m = re.search(r"\{.*\}", text, re.DOTALL)
    if not m:
        raise ValueError(f"Resume parser did not return JSON. Got: {text[:300]}")
    return json.loads(m.group(0))


async def parse_to_profile(resume_text: str) -> dict:
    """Run resume text through Claude and return a structured profile dict."""
    text = (resume_text or "").strip()
    if not text:
        raise ValueError("Empty resume text")
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + "\n\n[truncated]"

    options = ClaudeAgentOptions(
        system_prompt=_EXTRACT_SYSTEM,
        model=_model(),
        allowed_tools=[],
        max_turns=1,
    )
    raw_chunks: list[str] = []
    user_prompt = f"Resume text:\n\n{text}\n\nReturn the JSON profile."
    async for msg in query(prompt=user_prompt, options=options):
        if isinstance(msg, AssistantMessage):
            for block in msg.content:
                if isinstance(block, TextBlock):
                    raw_chunks.append(block.text)
    raw = "".join(raw_chunks)
    data = _extract_json(raw)

    # Sanitise nulls / missing fields
    return {
        "name": (data.get("name") or "").strip() or None,
        "headline": (data.get("headline") or "").strip(),
        "background": (data.get("background") or "").strip(),
        "role_context_suggestion": (data.get("role_context_suggestion") or "").strip(),
        "focus_areas_suggestion": (data.get("focus_areas_suggestion") or "").strip(),
        "topic_suggestion": (data.get("topic_suggestion") or "").strip(),
    }
